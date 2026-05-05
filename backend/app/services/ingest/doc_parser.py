import docx
import os
from typing import List, Dict, Any

class DocParser:
    def __init__(self):
        pass
    
    def parse(self, file_path):
        """
        解析Word文档，提取文本内容
        """
        blocks: List[Dict[str, Any]] = []
        chunk_index = 0
        try:
            doc = docx.Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text is None:
                    continue
                blocks.append(
                    {
                        "text": paragraph.text,
                        "page": None,
                        "chunk_index": chunk_index,
                    }
                )
                chunk_index += 1
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    cells = [(cell.text or "") for cell in row.cells]
                    row_text = "\t".join(cells) + "\t"
                    blocks.append(
                        {
                            "text": row_text,
                            "page": None,
                            "chunk_index": chunk_index,
                        }
                    )
                    chunk_index += 1
        except Exception as e:
            print(f"文档解析错误: {str(e)}")
        
        text = "\n".join([b.get("text", "") for b in blocks]) + ("\n" if blocks else "")
        return {
            "text": text,
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "num_paragraphs": len(doc.paragraphs) if 'doc' in locals() else 0,
                "num_tables": len(doc.tables) if 'doc' in locals() else 0
            },
            "blocks": blocks,
        }
